def df_to_word(student_id, student_name, open_ele_id, open_ele):
	from docx import Document
	from docx.enum.text import WD_ALIGN_PARAGRAPH
	from docx.enum.table import WD_TABLE_ALIGNMENT
	from docx.enum.table import WD_ALIGN_VERTICAL

	from docx.shared import Pt, Inches

	document = Document()

	p = document.add_paragraph()
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	run = p.add_run('Punjab Engineering College, Chandigarh\n(Deemed To Be University)\n\n')
	run.bold = True
	font = run.font
	font.name = 'Times New Roman'
	font.size = Pt(24)

	run = p.add_run('STUDENT REGISTRATION FORM\n\n')
	run.bold = True
	font = run.font
	font.name = 'Times New Roman'
	font.size = Pt(18)

	p = document.add_paragraph()
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

	p.add_run('Student ID: ')
	p.add_run('{}\t\t\t\t\t      '.format(student_id)).bold = True

	p.add_run('Name: ')
	p.add_run('{}\n'.format(student_name)).bold = True

	p.add_run('Department: CSE\t\t\t')
	p.add_run('Academic Session: 2020-21\t\t')
	p.add_run('Semester: 7th\n')


	records = (
				(1, open_ele_id, open_ele),
				(2, dec1_id, dec1),
				(3, dec2_id, dec2),
				(4, 'CSN441', 'Major Project')
			  )

	table = document.add_table(rows=1, cols=3, style='Table Grid')

	headers = table.rows[0].cells
	headers[0].text = 'Sr. No.'
	headers[1].text = 'Course ID'
	headers[2].text = 'Course Name'

	for cell in headers:
		cell.vertical_allignment = WD_ALIGN_VERTICAL.CENTER

	for sr_no, course_id, course_name in records:
	    tds = table.add_row().cells
	    tds[0].text = str(sr_no)
	    tds[1].text = course_id
	    tds[2].text = course_name

	for cell in table.columns[0].cells:
	    cell.width = Inches(0.4)

	for cell in table.columns[1].cells:
	    cell.width = Inches(3)

	for cell in table.columns[2].cells:
	    cell.width = Inches(3)

	p = document.add_paragraph()
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
	p.add_run('\n\n')

	table = document.add_table(rows=1, cols=2, style='Table Grid')

	headers = table.rows[0].cells
	p = headers[0].add_paragraph('Faculty Incharge')
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER
	p = headers[1].add_paragraph('Signature of Student')
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER

	for cell in table.rows[0].cells:
		cell.vertical_allignment = WD_ALIGN_VERTICAL.BOTTOM

	for row in table.rows:
	    row.height = Inches(1.2)


	document.save('{}.docx'.format(student_id))
